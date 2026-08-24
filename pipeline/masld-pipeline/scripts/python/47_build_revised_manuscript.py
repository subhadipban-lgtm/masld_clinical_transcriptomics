#!/usr/bin/env python3
"""Build the REVISED manuscript docx from the 2026-08-24 temporal draft.

Strategy: load the original paragraphs (extracted to /tmp/manuscript_temporal.txt),
keep verified content verbatim, replace corrupted/contradicted paragraphs by prefix
match, insert new Methods/Results/Discussion blocks from WS26-WS28, rebuild Tables
2/3/4 and add Tables 6-8 from authoritative CSVs, renumber figures, embed figures
(originals re-used from the unpacked docx media + new WS26-28 PNGs).

Every number in replaced/new text is sourced from a file in this repository; the
numbers ledger records the mismatches that motivated each replacement.
"""
import html
import os

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SRC_TXT = "/tmp/manuscript_temporal.txt"
MEDIA = "/tmp/msx/word/media"
OUT = "manuscript_assets/MASLD_Ferroptosis_Matrisome_Temporal_AcademicPaper_2026-08-24_REVISED.docx"

paras = [html.unescape(l) for l in open(SRC_TXT).read().split("\n")]
paras = [p.strip() for p in paras if p.strip()]

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
pf = style.paragraph_format
pf.line_spacing = 1.3
pf.space_after = Pt(6)

for s in ["Heading 1", "Heading 2", "Heading 3"]:
    st = doc.styles[s]
    st.font.name = "Times New Roman"
    st.font.color.rgb = RGBColor(0, 0, 0)


def p(text, bold=False, italic=False, size=None, align=None, style_name=None):
    par = doc.add_paragraph()
    if style_name:
        par.style = doc.styles[style_name]
    run = par.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if align:
        par.alignment = align
    return par


def caption(text):
    par = p(text, italic=True, size=9)
    par.paragraph_format.space_after = Pt(12)
    return par


def figure(path, cap, width=6.1):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(path, width=Inches(width))
    caption(cap)


def table(df, cap, colw=None, numfmt=None):
    caption(cap)
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for j, c in enumerate(df.columns):
        hdr[j].text = ""
        r = hdr[j].paragraphs[0].add_run(str(c))
        r.bold = True
        r.font.size = Pt(9)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = ""
            txt = numfmt(v) if numfmt else str(v)
            r = cells[j].paragraphs[0].add_run(str(txt))
            r.font.size = Pt(9)
    for rowx in t.rows:
        for c in rowx.cells:
            for pp in c.paragraphs:
                pp.paragraph_format.space_after = Pt(2)
                pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def find(prefix, start=0):
    for i in range(start, len(paras)):
        if paras[i].startswith(prefix):
            return i
    raise ValueError("not found: " + prefix[:60])


# ---------------- replacement / new text ----------------
TITLE = ("An early, suppressor-biased ferroptosis response precedes matrisome convergence "
         "in MASLD fibrosis: a multi-cohort transcriptional benchmarking study")

ABS_BA = ("Background & Aims: Fibrosis stage is the strongest histological determinant of "
          "liver-related outcomes in metabolic dysfunction-associated steatotic liver disease "
          "(MASLD). Multiple hepatic transcriptional signatures of stage have been published, "
          "but they are rarely validated in fully independent cohorts, almost never compared "
          "head-to-head under a common protocol, and have not been evaluated at the fibrosis "
          "thresholds that now govern treatment eligibility. We asked where along the F0–F4 "
          "continuum the ferroptosis programme first diverges, whether it precedes the "
          "matrisome convergence that dominates late-stage signatures, and whether any "
          "ferroptosis-derived signal is clinically discriminative at the treatment boundary.")

ABS_METH = ("Methods: Differential expression (voom-limma) between late (F3–F4, n = 89) and "
            "early (F0–F2, n = 260) fibrosis across 349 biopsy-staged patients from three "
            "cohorts. Signature scores were validated in two independent cohorts (Fujiwara "
            "n = 213, UCAM n = 58) under a locked normalisation protocol and benchmarked "
            "against five published panels at F0–2 vs F3–4, F0–1 vs F2–4 (treatment "
            "eligibility), and F2 vs F3. Stage-resolved analysis combined FerrDb driver and "
            "suppressor scoring, per-adjacent-stage-pair GSEA (limma-trend contrasts), paired "
            "second biopsies (n = 58), a cell-death pathway comparison, and effector-gene "
            "localisation in a 62,210-cell single-cell atlas across 10 liver donors.")

ABS_RES = ("Results: The signature (649 genes identified, 500 measurable) discriminated late "
           "from early fibrosis with AUROC 0.811 (Fujiwara) and 0.826 (UCAM); the 75-gene "
           "panel reached 0.847 and 0.769. Transcription factor regulon analysis identified "
           "10 TFs with significant stage-dependent activity (FDR < 0.05), including SPIB, "
           "which differentiated all stage pairs including F0 vs F1 (FDR = 0.018) and "
           "replicated in the larger GSE193066 cohort (p = 1.15 × 10⁻⁵). Per-stage-pair GSEA "
           "showed the ferroptosis suppressor programme diverging at the earliest transition "
           "(F0 vs F1, NES 1.63, padj < 10⁻⁴), but the driver-minus-suppressor balance "
           "shifted toward suppressors with stage (ρ = −0.238, p = 7.2 × 10⁻⁶) — an early, "
           "suppressor-biased response rather than a balanced poised state. Matrisome "
           "enrichment of discriminating genes emerged from rank 25 onward. A ferroptosis-DEG "
           "panel failed the treatment boundary in Fujiwara (AUROC 0.588 at F0–1 vs F2–4, CI "
           "including 0.50). Among cell-death programmes, necroptosis showed the strongest "
           "stage association (p = 2.7 × 10⁻¹⁰), exceeding ferroptosis (p = 5.1 × 10⁻⁹). In "
           "the single-cell atlas no pro-ferroptotic effector gene was enriched in hepatic "
           "stellate cells across donors; iron-import genes trended lower. TRRUST analysis "
           "identified ZBTB2 as a regulator of ferroptosis genes (fold = 10.62, p = 0.0013, "
           "FDR = 0.0066). All five published panels and our own were enriched for matrisome "
           "genes (fold 6.3–16.2, all padj < 0.05). Performance degraded for every panel at "
           "the clinically decisive thresholds (F0–1 vs F2–4: 0.770–0.786; F2 vs F3 rests on "
           "Fujiwara alone, 0.746). Ten rigorously evaluated null findings are reported.")

ABS_CONC = ("Conclusions: The ferroptosis signal in MASLD fibrosis is an early, "
            "suppressor-biased transcriptional response (F0→F1) that precedes matrisome "
            "convergence but does not itself discriminate at the F2 treatment boundary — "
            "a biological ceiling of bulk transcriptomics that is not rescued by advanced "
            "modelling or clinical covariate fusion. Ferroptosis is not the dominant "
            "cell-death programme transcriptionally (necroptosis outranks it), and its "
            "effector genes are not localised to stellate cells. Clinical deployment is not "
            "advocated on present evidence.")

KEYWORDS = ("MASLD; fibrosis staging; transcriptional signature; matrisome; ferroptosis; "
            "necroptosis; treatment eligibility; benchmarking; transcription factor regulon; "
            "temporal sequence")

INTRO_POISED = ("The transition from F0 to F1 represents the earliest detectable molecular "
                "shift in MASLD fibrosis, yet its drivers remain incompletely characterized. "
                "Recent evidence has implicated ferroptosis—an iron-dependent, non-apoptotic "
                "form of cell death driven by lipid peroxidation—as a critical mechanism in "
                "early metabolic injury (Peleman et al., 2024a, 2024b; Zhou et al., 2026). "
                "Experimental studies have demonstrated that inhibiting ferroptosis prevents "
                "progression of steatotic liver disease in obese mice (Park et al., 2024), and "
                "integrated hepatic ferroptosis gene signatures have been shown to dictate "
                "pathogenic features of ferroptosis in MASLD (Matsumoto et al., 2025). Because "
                "ferric cell-death programmes involve concurrently upregulated driver and "
                "suppressor arms (Peleman et al., 2024b; Wang, 2026), the balance between the "
                "two arms—not either arm alone—carries the biological interpretation: a "
                "balanced state would be consistent with cellular 'poising', whereas a "
                "biased state would indicate a directed response. Whether and when such a "
                "response precedes matrisome-driven transcriptional convergence has not been "
                "systematically evaluated.")

METHODS_COHORTS = ("Three bulk RNA-seq cohorts with verified histological staging were "
                   "harmonised into the discovery set: GSE135251 (n = 216), GSE130970 "
                   "(n = 78) and GSE185051 (paediatric, n = 55; 349 patients total). Two "
                   "additional GEO series (GSE167523, GSE126848) were excluded because their "
                   "binary NASH labels could not be mapped to histological stage. External "
                   "validation used the Fujiwara cohort (n = 213; SuperSeries GSE193084, "
                   "comprising sub-series GSE192959, GSE193066 and GSE193080; "
                   "Fujiwara-Contreras et al., 2022) and the UCAM cohort (n = 58). All "
                   "cohorts were biopsy-staged using the NASH CRN system (F0–F4) and were "
                   "verified to share no samples with the discovery set (Fig. S1). Age was "
                   "available for 133 of 349 discovery samples and sex for 133 (fields for "
                   "GSE135251 could not be parsed from the series metadata).")

METHODS_TF = ("Transcription factor (TF) regulon inference was performed using pySCENIC "
              "(Aibar et al., 2017) with AUCell scoring against the CollecTRI regulatory "
              "network on the Fujiwara validation matrix. The top 20 TFs by activity "
              "variance across samples were prioritized for downstream analysis. "
              "Differential TF activity was tested with Mann–Whitney U comparing Stage 3 "
              "(n = 30) versus Stage 0 (n = 12), with Benjamini–Hochrich correction across "
              "TFs; post-hoc pairwise Mann-Whitney U tests with FDR correction were applied "
              "for SPIB across all stage pairs, stratified by cohort. Single-cell "
              "localisation of TF expression used an independent atlas (Section 2.7.10).")

METHODS_NEW = [
    ("2.7.7 Ferroptosis programme across the fibrosis continuum",
     "Ferroptosis drivers (FerrDb V2 filtered, n = 264) and suppressors (n = 238) were "
     "scored in every sample of the locked discovery matrix as the mean within-sample "
     "z-score of the set members present (194 and 186 genes respectively). Stage association "
     "used Kruskal–Wallis across F0–F4 with Dunn's post-hoc tests (BH within family). A "
     "ferroptosis-DEG classifier was derived from signature DEGs that are FerrDb members "
     "(adj.P < 0.05, |log2FC| > 0.5; 9 genes) and scored as the mean z weighted by the sign "
     "of the discovery log2FC; AUROCs with class-stratified bootstrap 95% CIs (2,000 "
     "resamples, seed 42) were computed at both clinical thresholds in all three cohorts, "
     "with the discovery row labelled in-sample. Age and sex associations were tested "
     "within cohorts (Spearman; Mann–Whitney) with partial rank correlations controlling "
     "stage, because age is confounded with cohort."),
    ("2.7.8 Stage-resolved enrichment, paired biopsies, and balance",
     "For each adjacent stage pair (F0vF1, F1vF2, F2vF3, F3vF4) a differential-expression "
     "contrast was fitted on the locked log2-CPM discovery matrix using limma with "
     "eBayes(trend = TRUE) (limma-trend; the original pooled contrast used voom precision "
     "weights on raw counts, which are not re-derivable for the locked universe — this "
     "deviation is stated). Weighted running-sum GSEA (1,000 label permutations, seed 42) "
     "tested the two FerrDb sets against each pair's t-ranked list, BH across all eight "
     "tests. Paired second-biopsy expression (raw counts, 58 patients with both timepoints) "
     "was library-size normalised, Ensembl-mapped, log2-transformed and z-scored across all "
     "116 paired samples; timepoint order was verified empirically against the locked "
     "baseline scores (ρ = 0.910 with baseline vs 0.668 with follow-up). Trajectory groups "
     "were defined by Δfibrosis (>0 progressor, <0 regressor, =0 stable) and tested with "
     "Kruskal–Wallis and continuous Spearman. The driver-minus-suppressor balance (a "
     "difference, not a ratio — the mean-z denominator of a ratio crosses zero) was tested "
     "against stage with Spearman and Kruskal–Wallis."),
    ("2.7.9 Cell-death context and co-expression network",
     "Apoptosis (GO:0006915), autophagy (GO:0006914), necroptosis (GO:0097300), pyroptosis "
     "(GO:0141201) and ferroptosis (GO:0097707) gene sets were downloaded from QuickGO "
     "(descendants, human; MSigDB C5 was not used because its download endpoint is "
     "login-restricted) and scored with the same mean-z estimator; Kruskal–Wallis across "
     "stages with BH across the seven-pathway family. Co-expression network analysis "
     "correlated all 347 measurable FerrDb genes across the 349 discovery samples (edges at "
     "|Spearman ρ| > 0.5) against a 1,000-replicate null in which each gene's sample vector "
     "is independently permuted (seed 42); community structure used greedy modularity."),
    ("2.7.10 Single-cell effector-gene localisation",
     "The GSE136103 atlas (Ramachandran et al., 2020; 10 human liver donors across FACS "
     "fractions; blood and mouse samples excluded) was processed from raw 10x matrices: "
     "cells with library size > 500 were CPM-normalised, log1p-transformed, and assigned to "
     "eight canonical types by the maximum mean-z of fixed marker panels (hepatocyte, "
     "stellate/mesenchymal, LSEC, cholangiocyte, macrophage, T/NK, B, erythroid; 62,210 "
     "cells). Fourteen ferroptosis effector genes (GPX4, ACSL4, SLC7A11, SLC3A2, TFRC, "
     "NFE2L2, SAT1, LPCAT3, ALOX15, PRNP, DHODH, CISD1, SLC7A5, STEAP3) were tested "
     "hepatocyte-versus-stellate by Wilcoxon signed-rank across donor-level pseudobulk "
     "means (cells were never treated as replicates), BH across genes."),
]

RES_32_TF = ("Stage-dependent TF activity. Mann–Whitney U testing of regulon activity "
             "(Stage 3, n = 30, versus Stage 0, n = 12, in the Fujiwara matrix) identified "
             "10 TFs at FDR < 0.05 (Table 3). Nine of the ten showed higher mean activity "
             "at Stage 3 (ARID5B alone decreased). TOX3 showed the most significant "
             "difference (p = 2.75 × 10⁻⁵, FDR = 0.0073) and ARHGAP35 the largest increase "
             "(Δ = +0.064, FDR = 0.0161).")

RES_32_SC = ("Single-cell localization of active TFs. Cell-type mean expression of the "
             "prioritised TFs in an independent atlas showed highest stellate-lineage "
             "expression for EID1 (mean 1.03 in stellate/fibroblast types vs 0.59 in "
             "hepatocytes) and NUCKS1 (0.87 vs 0.82), with E2F7 and ZBTB7C directionally "
             "higher in stellate cells (log2 ratios 1.9 and 0.5) but not significantly so "
             "in donor-level testing (Section 3.17). These observations anchor the bulk "
             "regulon signal to ECM-producing lineages without establishing stellate-cell "
             "specificity.")

RES_33 = ("To establish TRRUST-documented regulatory relationships, TF–target interactions "
          "of the 20 prioritised TFs were extracted from TRRUST v2 with recorded Mode and "
          "PMID evidence. This yielded 10 interactions across 5 TFs (Table 4); only "
          "ZBTB2's edges target FerrDb ferroptosis genes (3 of its 4 targets: CDKN1A, "
          "MDM2, TP53). Hypergeometric enrichment of ferroptosis genes among ZBTB2 targets "
          "was significant (fold = 10.62, p = 0.0013, FDR = 0.0066), identifying a "
          "TRRUST-documented ZBTB2–TP53/MDM2 axis connected to ferroptosis regulation "
          "(Fig. 4).")

RES_361 = ("Early fusion of transcriptomic signatures with standard clinical covariates "
           "(age, BMI, diabetes status) failed to improve discrimination at the treatment "
           "boundary. In five model architectures (logistic regression, random forest, "
           "XGBoost, LightGBM, multilayer perceptron), transcriptomic-only models "
           "outperformed fusion models in every case, suggesting that clinical covariates "
           "contribute minimal additional information at this boundary.")

RES_310 = ("In 58 paired biopsies, the signature score did not track histological "
           "regression. After adjustment for technical covariates (log library size, genes "
           "detected, housekeeping median), the follow-up timepoint effect was attenuated "
           "from p = 0.005 to p = 0.058, and the responder × timepoint interaction was null "
           "(p = 0.99); the continuous trajectory therefore does not detect dynamic "
           "therapeutic regression.")

NEW_RESULTS = []

# ---------------- compose ----------------
i_title = find("Ferroptosis poising precedes")
p(TITLE, bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
for pref in ["[Author Name]", "[Affiliation]", "Corresponding author", "Email:"]:
    for j in [x for x in paras if x.startswith(pref)]:
        p(j, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

p("Abstract", bold=True, size=12)
p(ABS_BA); p(ABS_METH); p(ABS_RES); p(ABS_CONC)
p("Keywords: " + KEYWORDS, italic=True)

# Introduction
doc.add_heading("1. Introduction", level=1)
intro_start = find("Metabolic dysfunction-associated steatotic liver disease (MASLD) is now")
p(paras[intro_start])
p(INTRO_POISED)
for pref in ["At later stages, fibrosis progression", "Integrative multi-omics studies",
             "While multiple hepatic transcriptional signatures"]:
    p(paras[find(pref)])
p(paras[find("We addressed these gaps directly.")].replace(
    "and performed transcription factor regulon analysis to identify early regulators of "
    "the ferroptosis-matrisome transition",
    "characterised the stage-resolved emergence of the ferroptosis programme against the "
    "matrisome convergence, and performed transcription factor regulon analysis to identify "
    "early regulators"))

# Methods
doc.add_heading("2. Methods", level=1)
doc.add_heading("2.1 Cohorts", level=2)
p(METHODS_COHORTS)
doc.add_heading("2.2 Processing", level=2)
p(paras[find("Raw counts were filtered")])
doc.add_heading("2.3 Differential expression", level=2)
p(paras[find("voom precision weights + lmFit")])
doc.add_heading("2.4 Signature scoring and thresholds", level=2)
p(paras[find("Each gene was z-scored")])
doc.add_heading("2.5 Comparator panels", level=2)
p("Five published fibrosis-stage panels from the Kamzolas et al. (2026) data release were "
  "benchmarked under identical normalisation, z-scoring and directionality rules: the "
  "3-gene external validation panel, the 15-gene ELBOW panel, the 57-biomarker panel, the "
  "145-gene variable-importance panel, and the 194-gene proteo-transcriptomic signature. "
  "Resolved gene counts and match rates per panel are reported in the load manifest; the "
  "194-gene panel resolved 174/193 symbols (91.6%) after explicit alias handling.")
doc.add_heading("2.6 Panel-size selection", level=2)
p(paras[find("The pre-specified rule was")])
doc.add_heading("2.7 Exploratory data modelling", level=2)
for h, pref in [("2.7.1 Pseudotime Trajectory Mapping", "PCA was fitted on baseline"),
                ("2.7.2 Matrisome enrichment analysis", "Hypergeometric tests evaluated"),
                ("2.7.3 Ferroptosis enrichment", "Gene-set enrichment of the signed-t"),
                ("2.7.4 Transcription factor regulon analysis", "PLACEHOLDER_TF"),
                ("2.7.5 TRRUST network analysis", "TF-target interactions were filtered to those"),
                ("2.7.6 Knowledge graph and drug–gene interactions", "A knowledge graph was constructed")]:
    doc.add_heading(h, level=3)
    if "PLACEHOLDER" in pref:
        p(METHODS_TF)
    else:
        p(paras[find(pref)])
for h, body in METHODS_NEW:
    doc.add_heading(h, level=3)
    p(body)
doc.add_heading("2.8 Sensitivity Analysis and Signature Stability", level=2)
p(paras[find("DEG identification was re-run")])
doc.add_heading("2.9 Bulk deconvolution", level=2)
p(paras[find("Cell-type associations were assessed by non-negative")])
doc.add_heading("2.10 Statistics", level=2)
p(paras[find("AUROCs with 95% confidence intervals")])
p("Bootstrap 95% confidence intervals (2,000 class-stratified patient-level resamples, "
  "seed 42) accompany every AUROC introduced in Sections 2.7.7–2.7.10. Stochastic analyses "
  "report the seed and permutation count used.")

# Results
doc.add_heading("3. Results", level=1)
doc.add_heading("3.1 Discovery of the 649-gene fibrosis signature", level=2)
p(paras[find("Differential expression between late (F3–F4, n = 89)")])
t1 = pd.DataFrame({"Characteristic": ["Total patients (discovery)", "Late fibrosis (F3–F4)",
                                      "Early fibrosis (F0–F2)", "Validation: Fujiwara (GSE193084)",
                                      "Validation: UCAM", "Genes tested (voom)",
                                      "Signature genes identified",
                                      "Measurable in validation universe"],
                   "Value": [349, 89, 260, 213, 58, "15,223", 649, "500 (77.0%)"]})
table(t1, "Table 1. Discovery cohort characteristics.")
figure(f"{MEDIA}/994f01431ec1dd207157fc02d578b510e9683ebf.png",
       "Figure 1. Discovery differential expression and matrisome enrichment. (A) Volcano "
       "plot of Late (F3–F4, n = 89) versus Early (F0–F2, n = 260); coloured points are the "
       "460 up-regulated (red) and 189 down-regulated (blue) genes at adj.P < 0.05 and "
       "|log2FC| > 0.5; ten most significant genes labelled. Matrisome membership is absent "
       "from the top ten discriminating genes (2 of 10, p = 0.061) and emerges from rank 25 "
       "onward (8 of 25, p = 4.5 × 10⁻⁶), consistent with matrisome convergence as a "
       "late-stage phenomenon (source: matrisome_by_rank.csv). (B) NABA matrisome enrichment "
       "fold across the ten canonical categories (NABA_MATRISOME fold 4.68, padj = 1.74 × "
       "10⁻⁴⁹).")

doc.add_heading("3.2 Transcription factor regulon analysis reveals stage-specific regulatory shifts", level=2)
p(paras[find("To identify upstream regulators of the fibrotic program")])
p(RES_32_TF)
d3 = pd.read_csv("results/ws19/diff_activity_S3_vs_S0_sig.csv")
d3 = d3.rename(columns={"Mean_Stage0": "Stage-0 mean", "Mean_Stage3": "Stage-3 mean",
                        "Mean_Diff (S3-S0)": "Δ (S3−S0)", "P_value": "p (MWU)",
                        "FDR_adj_P": "FDR"})
d3 = d3[["TF", "Stage-0 mean", "Stage-3 mean", "Δ (S3−S0)", "p (MWU)", "FDR"]]
table(d3, "Table 3. Differentially active transcription factors (Mann–Whitney U, Stage 3 "
          "vs Stage 0, n = 30/12; BH across 1,099 tested regulons). Source: "
          "diff_activity_S3_vs_S0_sig.csv.", numfmt=lambda v: f"{v:.4g}" if isinstance(v, float) else v)
p(paras[find("SPIB shows consistent stage-dependent activity")])
p(RES_32_SC)
figure(f"{MEDIA}/9fe538c56a0f6617b1bb582abda7db89d731e0ce.png",
       "Figure 3. TF regulon analysis: patient clustermap of top 20 transcription factors.")
figure(f"{MEDIA}/bd8bbd07a608f339a6fc0d1bda74ab5de32edc2e.png",
       "Figure 3B. Single-cell localization of TF expression across liver cell lineages "
       "(cell-type mean expression; inferential donor-level testing in Section 3.17).")
figure(f"{MEDIA}/371e931d1b0c3cc4a45425c5f9a2814d88d200c7.png",
       "Figure 3C. SPIB activity by fibrosis stage and cohort. Left: GSE193066 (n = 106, "
       "Kruskal-Wallis p = 1.15 × 10⁻⁵). Right: GSE192959 (n = 42, p = 0.288). SPIB "
       "differentiates all stage pairs including F0 vs F1 (FDR = 0.018).")

doc.add_heading("3.3 TRRUST-documented regulatory network identifies ZBTB2 as a ferroptosis regulator", level=2)
p(RES_33)
t4 = pd.read_csv("/tmp/trrust_table4_full.csv")
t4 = t4.rename(columns={"ferro_target": "Ferroptosis target"})
t4["Ferroptosis target"] = t4["Ferroptosis target"].map({True: "Yes", False: "No"})
table(t4, "Table 4. TRRUST-documented interactions of the 20 prioritised TFs (all edges "
          "with recorded Mode and PMID; rebuilt from trrust_rawdata.human.tsv). Only "
          "ZBTB2's edges target FerrDb ferroptosis genes.")
figure(f"{MEDIA}/924e02737fc452cd5004b95531f650f5b7994569.png",
       "Figure 4. ZBTB2 regulatory axis connected to ferroptosis genes (TRRUST-documented "
       "modes; fold 10.62, p = 0.0013, FDR = 0.0066).")

doc.add_heading("3.4 Signature convergence with an independently derived panel", level=2)
p("Testing the 649-gene signature for membership of the Kamzolas 145 panel showed "
  "substantial convergence: 104 of 649 signature genes (16.0%) were shared (expected 5.7, "
  "fold enrichment = 18.34, p = 4.07 × 10⁻¹¹⁸), confirming that independently derived MASLD "
  "fibrosis signatures overlap heavily. GSEA of the genome-wide ranked list showed a trend "
  "toward enrichment of ferroptosis suppressor genes in advanced fibrosis (NES = 1.264, "
  "p = 0.065, FDR = 0.13) that did not survive correction.")

doc.add_heading("3.5 External validation and panel benchmarking", level=2)
p("The 649-gene signature (500 measurable genes) was validated in two independent cohorts "
  "under a locked normalisation protocol. At the conventional F0–2 vs F3–4 threshold, the "
  "full signature achieved AUROC 0.811 (Fujiwara) and 0.826 (UCAM), with the 75-gene panel "
  "at 0.847 and 0.769; both outperformed all five published comparator panels in Fujiwara, "
  "while in UCAM formal non-inferiority was not met against the in-sample 15-gene ELBOW "
  "panel (Table 2; Fig. 2).")
t3f = pd.read_csv("results/ws15/table3_final.csv")
piv = t3f[~t3f.apparent].pivot_table(index=["split", "cohort"], columns="panel", values="AUROC")
piv = piv.reset_index()
piv.columns.name = None
order = ["Ours_full", "Ours_75", "Ours_10", "Kamzolas_145", "Kamzolas_57BM",
         "Kamzolas_15BM", "Kamzolas_194PT", "Kamzolas_3gene"]
piv = piv[["split", "cohort"] + order]
piv = piv.rename(columns={"Ours_full": "Ours-649", "Ours_75": "Ours-75", "Ours_10": "Ours-10",
                          "Kamzolas_145": "Kamz-145", "Kamzolas_57BM": "Kamz-57BM",
                          "Kamzolas_15BM": "Kamz-15", "Kamzolas_194PT": "Kamz-194",
                          "Kamzolas_3gene": "Kamz-3"})
table(piv, "Table 2. External validation AUROCs (held-out strata; rebuilt from "
          "table3_final.csv, WS15 locked build — the comparator set is the five "
          "Kamzolas-release panels).", numfmt=lambda v: f"{v:.3f}" if isinstance(v, float) else v)
figure(f"{MEDIA}/1ea1e928892ed0fae6c235e23e77937325a30646.png",
       "Figure 2. External validation and panel benchmarking at three clinically defined thresholds.")

doc.add_heading("3.6 Performance at clinically defined thresholds", level=2)
p("Performance degraded for every panel at the clinically decisive thresholds. At F0–1 vs "
  "F2–4 (treatment eligibility), the 75-gene panel achieved AUROC 0.770 in Fujiwara and "
  "0.786 in UCAM. At F2 vs F3 the evidence rests on Fujiwara alone (75-gene AUROC 0.746; "
  "full signature 0.733) because the UCAM split was underpowered (n = 12 F2, n = 7 F3). "
  "Under 5-fold cross-validation, transcriptomic-only models outperformed early fusion in "
  "five of five architectures on AUROC; the best transcriptomic model reached AUROC 0.755 "
  "(random forest) and 0.753 (logistic regression), against 0.719 and 0.694 for the "
  "corresponding fusion models.")
figure(f"{MEDIA}/db9f72f8eac877da53e74ad0a4f448444c1e88af.png",
       "Figure 6. Prioritisation and panel characterisation at treatment thresholds.")
doc.add_heading("3.6.1 Multi-omic fusion performance", level=2)
p(RES_361)

doc.add_heading("3.7 Panel size is not identifiable from performance data", level=2)
p(paras[find("Ranking the measurable signature by discovery |t|")])
p(paras[find("The decisive observation is not the ordering")])
p(paras[find("We therefore report that panel size is not identifiable")])
figure(f"{MEDIA}/c17b01ca276aa4a65f25990056aadd36669c413d.png",
       "Figure 5. Panel size parsimony analysis.")

doc.add_heading("3.8 Fixed threshold is only partially transportable", level=2)
p(paras[find("A fixed operating point derived")])

doc.add_heading("3.9 Comparison with clinical scores: parity with FIB-4", level=2)
p(paras[find("In a 48-patient clinical subset")])
t5 = pd.DataFrame({"Score": ["FIB-4", "APRI", "Ours (5-gene)"],
                   "AUROC": [0.555, 0.592, 0.561],
                   "95% CI": ["0.366–0.732", "0.416–0.756", "0.375–0.733"],
                   "n": [47, 47, 47]})
table(t5, "Table 5. Clinical score comparison (48-patient matched subset; DeLong p > 0.35 for all pairs).")
figure(f"{MEDIA}/339a7d88ce3b893e796a3262d529385b7d8dced4.png",
       "Figure 7. Benchmarking against clinical scores in the matched subset.")

doc.add_heading("3.10 Longitudinal analysis: technical artefact", level=2)
p(RES_310)
figure(f"{MEDIA}/6f45f425a57b24b9b8254e23e1bb16cb23137eaa.png",
       "Figure 8. Longitudinal pseudotime analysis.")

doc.add_heading("3.11 Cell-type localisation: compositional correction required", level=2)
p(paras[find("Deconvolution revealed a cholangiocyte-dominant association")])

doc.add_heading("3.12 Exploratory analyses", level=2)
doc.add_heading("3.12.1 WGCNA", level=3)
p(paras[find("A properly re-run weighted co-expression analysis")])
doc.add_heading("3.12.2 Ferroptosis", level=3)
p(paras[find("Gene-set enrichment of the signed-t ranking")])
doc.add_heading("3.12.3 Sex", level=3)
p(paras[find("Clean sex labels were available")])
doc.add_heading("3.12.4 Extra-hepatic pathways", level=3)
p(paras[find("Only REACTOME_HEMOSTASIS was enriched")])

# NEW RESULTS 3.13-3.17
doc.add_heading("3.13 The ferroptosis programme across the fibrosis continuum", level=2)
p("Across the 349 discovery samples, the suppressor-set score increased with stage "
  "(Kruskal–Wallis p = 1.2 × 10⁻⁶; six of ten stage pairs at FDR < 0.05, all involving F0), "
  "whereas the driver-set score was only weakly associated (p = 0.031; two of ten pairs) — "
  "an asymmetric, suppressor-led response (Fig. 9). A nine-gene ferroptosis-DEG classifier "
  "(two drivers, seven suppressors) discriminated F3+ from F0–2 in Fujiwara at AUROC 0.659 "
  "(95% CI 0.589–0.734) but failed the treatment boundary at 0.588 (0.499–0.673, CI "
  "including 0.50); UCAM gave 0.732 and 0.782 with wide intervals (n = 58) (Table 6). "
  "Ferroptosis genes alone therefore do not substitute for the matrisome-dominated "
  "signature at the clinically decisive threshold. Age and sex associations were tested "
  "within cohorts: the pooled age correlation (ρ = −0.58/−0.72) was cohort-confounded "
  "(age missing for all GSE135251 samples; the pooled set spans adult and paediatric "
  "cohorts) and was null within the adult cohort (GSE130970: ρ = −0.21 and −0.07); sex "
  "showed no association in cohorts with valid labels (all p > 0.36).")
figure("results/ws26_ferroptosis_additional/fig_2_2_gsva_boxplots.png",
       "Figure 9. Ferroptosis driver (194 genes) and suppressor (186 genes) mean-z scores "
       "across fibrosis stages, discovery 349. Kruskal–Wallis p = 0.031 (drivers) and "
       "1.2 × 10⁻⁶ (suppressors); Dunn's post-hoc with BH.")
t6 = pd.read_csv("results/ws26_ferroptosis_additional/ferroptosis_signature_auroc.csv")
t6 = t6[["cohort", "threshold", "AUROC", "CI95_lo", "CI95_hi", "n", "n_positive",
         "genes_measured", "in_sample_discovery"]]
t6 = t6.rename(columns={"CI95_lo": "CI lo", "CI95_hi": "CI hi", "n_positive": "n pos",
                        "genes_measured": "genes", "in_sample_discovery": "in-sample"})
t6["in-sample"] = t6["in-sample"].map({True: "Yes", False: "No"})
table(t6, "Table 6. Ferroptosis-DEG classifier (9 genes, discovery-weighted signs): AUROC "
          "with bootstrap 95% CIs. Source: ferroptosis_signature_auroc.csv.",
      numfmt=lambda v: f"{v:.3f}" if isinstance(v, float) else v)

doc.add_heading("3.14 Stage-resolved emergence: the suppressor programme diverges at F0→F1", level=2)
p("Per-adjacent-stage-pair GSEA located the divergence point. The suppressor set was "
  "enriched among genes up-regulated at the earliest transition (F0 vs F1: NES 1.63, BH "
  "padj < 10⁻⁴) and again at F1 vs F2 (NES 1.48, padj = 0.013), with no signal at F2 vs F3 "
  "or F3 vs F4 (the latter low-powered, n = 73/16). The driver set showed a weaker "
  "early-pair signal (NES 1.47, padj = 0.024) and nothing thereafter (Fig. 10). The "
  "ferroptosis transcriptional response is therefore an early-fibrosis event that "
  "precedes—and is exhausted before—the matrisome-dominated programme, which becomes "
  "apparent in the discriminating-gene ranking only from rank 25 onward (Section 3.1).")
figure("results/ws27_ferroptosis_followups/fig_stagepair_gsea.png",
       "Figure 10. Ferroptosis set enrichment per adjacent stage transition (limma-trend "
       "contrasts; BH across 8 tests; * padj < 0.05). Source: stagepair_gsea_results.csv.")

doc.add_heading("3.15 Paired biopsies: ferroptosis scores do not track fibrosis trajectory", level=2)
p("In the 58 patients with paired biopsies (timepoint order verified against locked "
  "baseline scores, ρ = 0.910 vs 0.668), the change in neither driver nor suppressor score "
  "differed across trajectory groups (15 progressors, 15 regressors, 28 stable; "
  "Kruskal–Wallis p = 0.76 and 0.82) nor correlated continuously with Δfibrosis (ρ = 0.09, "
  "p = 0.52; ρ = 0.01, p = 0.95). Both scores decreased globally at follow-up (mean Δ "
  "−0.27 and −0.31; Wilcoxon p ≈ 2–3 × 10⁻⁵), an unadjusted timepoint effect consistent "
  "with the technical artefact characterised in Section 3.10 (Fig. 11).")
figure("results/ws27_ferroptosis_followups/fig_paired_trajectory.png",
       "Figure 11. Ferroptosis score change (follow-up − baseline) by fibrosis trajectory, "
       "n = 58. Source: paired_ferroptosis_scores.csv / paired_trajectory_results.csv.")

doc.add_heading("3.16 Balance, cell-death context, and network structure", level=2)
p("The driver-minus-suppressor balance declined with stage (Spearman ρ = −0.238, "
  "p = 7.2 × 10⁻⁶; Kruskal–Wallis p = 2.3 × 10⁻⁴, n = 349): the response is "
  "suppressor-biased and becomes more so as fibrosis advances, which contradicts a "
  "'balanced poising' interpretation. Among cell-death programmes scored identically, "
  "necroptosis showed the strongest stage association (Kruskal–Wallis p = 2.7 × 10⁻¹⁰, "
  "ρ = +0.374), exceeding ferroptosis (GO:0097707: p = 5.1 × 10⁻⁹, ρ = +0.353); apoptosis, "
  "autophagy and pyroptosis were also positively associated (Table 7; Fig. 12) — all death "
  "programmes co-move with stage and none is uniquely 'dominant'. Finally, co-expression "
  "network analysis of the 347 measurable ferroptosis genes produced 14,917 edges at "
  "|ρ| > 0.5 against a per-gene permutation null of zero, but with low modularity "
  "(Q = 0.13; a quarter of all pairs wired): drivers and suppressors are extensively "
  "co-expressed with each other (cross-type edge fraction 0.55 vs 0.46 expected, p ≈ "
  "10⁻¹⁶⁶), consistent with coordinated regulation rather than opposing arms.")
figure("results/ws28_celldeath/fig_celldeath_comparison.png",
       "Figure 12. Stage association of cell-death pathway scores (discovery 349; QuickGO "
       "sets). Source: celldeath_stage_association.csv.")
t7 = pd.read_csv("results/ws28_celldeath/celldeath_stage_association.csv")
t7 = t7[["pathway", "genes_used", "KW_p", "KW_p_adj", "spearman_rho_vs_stage"]]
t7 = t7.rename(columns={"genes_used": "genes", "KW_p": "KW p", "KW_p_adj": "KW padj",
                        "spearman_rho_vs_stage": "ρ vs stage"})
table(t7, "Table 7. Cell-death pathway stage association (BH across 7). Source: "
          "celldeath_stage_association.csv.",
      numfmt=lambda v: f"{v:.2e}" if isinstance(v, float) and v < 0.01 else (f"{v:.3f}" if isinstance(v, float) else v))

doc.add_heading("3.17 Effector-gene localisation: no stellate-cell enrichment", level=2)
p("In the GSE136103 atlas (62,210 cells, 10 liver donors, marker-assigned types), no "
  "pro-ferroptotic effector gene was significantly enriched in hepatic stellate cells "
  "relative to hepatocytes in donor-level pseudobulk testing (Wilcoxon signed-rank across "
  "10 donors, BH across 14 genes; Table 8; Fig. 13). The only significant difference was "
  "lower STEAP3 in stellate cells (log2 ratio −2.29, padj = 0.027), with TFRC trending the "
  "same way (−0.43, padj = 0.23) — iron-import machinery is lower, not higher, in the "
  "stellate compartment. GPX4 (+0.14), SLC7A11 (+0.08) and ACSL4 (−0.27) were "
  "indistinguishable between the two lineages. The bulk-level ferroptosis signal is "
  "therefore not attributable to a stellate-cell effector programme.")
figure("results/ws28_sc_effectors/fig_effector_celltype_heatmap.png",
       "Figure 13. Ferroptosis effector genes by cell type (GSE136103; descriptive "
       "cell-level means — inferential tests use donor pseudobulk, Table 8).")
t8 = pd.read_csv("results/ws28_sc_effectors/effector_HSC_vs_hepatocyte.csv")
t8 = t8.rename(columns={"log2_HSC_over_hep": "log2 HSC/hep", "wilcoxon_p": "p",
                        "p_adj": "padj", "n_donors": "donors"})
t8 = t8[["gene", "HSC_mean", "hep_mean", "log2 HSC/hep", "p", "padj", "donors"]]
table(t8, "Table 8. Effector genes, stellate vs hepatocyte donor pseudobulk (n = 10 "
          "donors; BH across genes). Source: effector_HSC_vs_hepatocyte.csv.",
      numfmt=lambda v: f"{v:.3g}" if isinstance(v, float) else v)

# Discussion
doc.add_heading("4. Discussion", level=1)
doc.add_heading("An early, suppressor-biased ferroptosis response", level=2)
p("Three stage-resolved analyses converge on one picture. The ferroptosis suppressor "
  "programme diverges at the earliest transition (F0 vs F1, NES 1.63) and is exhausted by "
  "F2; the driver arm barely tracks stage; and the driver-minus-suppressor balance shifts "
  "toward suppressors as fibrosis advances (ρ = −0.238, p = 7.2 × 10⁻⁶). SPIB, the one "
  "regulon that differentiates every stage pair including F0 vs F1 (FDR = 0.018), sits at "
  "the front of this early response. Together these findings replace the 'balanced "
  "poising' hypothesis (Peleman et al., 2024b; Wang, 2026) with a directed one: early "
  "MASLD liver mounts a predominantly anti-ferroptotic transcriptional response, "
  "consistent with an adaptive defence against lipotoxic iron injury (Theys et al., 2024; "
  "Liu et al., 2025) rather than with cells sitting at a balanced point between death and "
  "survival.")
doc.add_heading("Matrisome convergence as the dominant late-stage programme", level=2)
p(paras[find("While all five published comparator panels")].replace(
    "Ferroptosis is", "Ferroptosis is") if False else
  "All five published comparator panels and our 649-gene signature were enriched for "
  "matrisome genes (fold 6.3–16.2, all padj < 0.05), yet matrisome membership was absent "
  "from the ten most discriminating genes (2 of 10, p = 0.061) and emerged only from rank "
  "25 onward (8 of 25, p = 4.5 × 10⁻⁶). The top-ranked genes instead reflect ductular "
  "reaction (STMN2, PDZK1IP1) and matrix crosslinking (LOXL4, FBLN5). This ordering — an "
  "early suppressor-biased ferroptosis response, then ductular and crosslinking genes, "
  "then broad matrisome convergence — is the temporal signature of a fibrotic liver "
  "mounting a defence and losing. The convergence itself is robust across derivation "
  "methods and aligns with the mechanosensitive feed-forward loop of matrisome remodelling: "
  "LOX-mediated collagen crosslinking stiffens the matrix, integrin receptors sense the "
  "changing physical force, and YAP/TAZ activation perpetuates stellate-cell activation "
  "and matrix deposition (Rudolph & Chin, 2024; Zhao et al., 2025; Fan et al., 2024; "
  "Benavides et al., 2026; Mayorca-Guiliani et al., 2025; Kołakowski et al., 2022; "
  "Steinberg et al., 2025).")
doc.add_heading("Ferroptosis in context: not dominant, not stellate-localised", level=2)
p("Two contextual results bound the ferroptosis interpretation. First, among five "
  "cell-death programmes scored identically, necroptosis showed the strongest stage "
  "association (p = 2.7 × 10⁻¹⁰) with ferroptosis second (5.1 × 10⁻⁹) and apoptosis, "
  "autophagy and pyroptosis all positive — bulk transcriptomics sees a general "
  "cell-death/stress programme riding with fibrosis, not a ferroptosis-specific cascade. "
  "Second, in 62,210 atlas cells across 10 donors, no pro-ferroptotic effector gene was "
  "enriched in stellate cells; iron-import genes (STEAP3 significantly, TFRC as a trend) "
  "were lower there. The frequently invoked 'ferroptosis-poised stellate cell' is not "
  "supported at the effector level in human liver. What survives is specific and modest: "
  "the ZBTB2–TP53/MDM2 TRRUST-documented axis (fold 10.62, FDR 0.0066), an early "
  "suppressor-set divergence, and a GSEA trend (NES 1.264, FDR 0.13) — hypothesis-generating "
  "observations, not an established convergent feature.")
doc.add_heading("A temporal framework for MASLD fibrosis", level=2)
p("The verified sequence is: (i) at F0→F1, a suppressor-biased ferroptosis response and "
  "SPIB regulon activity; (ii) through F1–F2, ductular-reaction and matrix-crosslinking "
  "genes dominate the discrimination ranking while the ferroptosis set signal fades; "
  "(iii) beyond F2, broad matrisome convergence that all published panels share. This "
  "ordering is consistent with early oxidative and lipotoxic stress engaging "
  "iron-dependent death defences (Termite et al., 2025; Gu et al., 2025; Carvalho et al., "
  "2024), and with phytochemical and experimental modulation of ferroptosis altering "
  "steatotic progression in mice (Park et al., 2024; Zhang et al., 2024; Zheng et al., "
  "2021). It does not, however, license ferroptosis-targeted staging claims: the "
  "ferroptosis-DEG classifier failed the F2 boundary (0.588, CI including 0.50) and "
  "paired-biopsy ferroptosis scores did not track regression.")
doc.add_heading("At the F2 treatment threshold, the transcriptional signal reaches a biological ceiling", level=2)
p("Despite the temporal insight, all signatures fail to resolve the F1→F2 transition. At "
  "F0–1 vs F2–4 the best-performing panel achieved an AUROC of 0.786, and at F2 vs F3 the "
  "evidence rests on a single cohort (0.746). Multiple independently derived signatures "
  "converge on the same matrisome programme and all fail at these boundaries, suggesting "
  "the molecular distinction between stages is not robustly captured by tissue-level gene "
  "expression. Early intervention targeting the ferroptosis-linked stress response may be "
  "most relevant before F2, while matrisome-directed therapies may be more appropriate "
  "later (Schwabe et al., 2025; Patel, 2025; Wang, 2026). Alternative approaches—single-"
  "cell resolution, spatial transcriptomics, or protein-based biomarkers—may be required "
  "at the eligibility boundary (Bong & Seo, 2026; Yang et al., 2026; Ko et al., 2026).")
p(paras[find("Discrimination was flat across panel sizes from 5 to 500")])
p(paras[find("The deconvolution analysis revealed a cholangiocyte-dominant")])
doc.add_heading("TRRUST-documented regulatory network links ZBTB2 to ferroptosis regulation", level=2)
p("The TRRUST network analysis identified ZBTB2 as a regulator of ferroptosis-related "
  "genes (fold = 10.62, p = 0.0013, FDR = 0.0066), with documented repression of TP53 and "
  "activation of MDM2 (PMID 19380588). GSEA showed a suppressor-set trend in advanced "
  "fibrosis (NES 1.264, FDR 0.13), and 104 of 649 signature genes are shared with the "
  "Kamzolas 145 panel (fold 18.34). The signature itself shows no ferroptosis enrichment; "
  "we therefore present these as hypothesis-generating regulatory observations.")
doc.add_heading("Matrisome remodelling, mechanotransduction, and matrikines", level=2)
p(paras[find("The convergent matrisome programme identified here")])
doc.add_heading("Ten rigorously evaluated null results", level=2)
p("The signature added no prognostic information beyond baseline stage in 58 paired "
  "biopsies. The fixed discovery-derived threshold transported only partially. Multi-omic "
  "fusion failed to improve treatment-boundary discrimination in all five architectures. "
  "A continuous pseudotime trajectory did not detect therapeutic regression, and its "
  "timepoint effect was a technical artefact after covariate adjustment. A re-run WGCNA "
  "found no module tracking stage. Ferroptosis genes were not enriched in the signature, "
  "did not discriminate at F2 as a standalone panel, and did not track paired-biopsy "
  "trajectory. No pro-ferroptotic effector gene localised to stellate cells. A sex×stage "
  "interaction did not survive restriction to adult cohorts.")

doc.add_heading("5. Conclusions", level=1)
p("Hepatic transcriptional signatures of MASLD fibrosis reveal a temporal sequence: an "
  "early, suppressor-biased ferroptosis response (F0→F1) — captured in part by the SPIB "
  "regulon — precedes the matrisome convergence that dominates the transcriptional "
  "landscape at later stages. The response is directed rather than balanced: the "
  "driver-minus-suppressor ratio shifts toward suppressors with stage. Ferroptosis is "
  "nevertheless not the dominant transcriptional cell-death programme (necroptosis "
  "outranks it), its effector genes are not stellate-localised, and a ferroptosis-only "
  "panel fails the treatment boundary. All signatures fail at F1→F2 — a biological ceiling "
  "of bulk transcriptomics that advanced machine learning and clinical covariate fusion "
  "do not rescue.")
p("Methodologically, we introduce a compositional correction for deconvolution-based "
  "cell-type claims, an explicitly nulled co-expression framework for gene-set coupling, "
  "and donor-level pseudobulk standards for atlas-level claims. The practical implication "
  "is clear: bulk transcriptomic staging panels currently hit a predictive ceiling at the "
  "most consequential clinical boundary, and their clinical deployment is not advocated "
  "on present evidence.")
p("The field should focus on single-cell resolution, spatial transcriptomics, or "
  "protein-based biomarkers to resolve the F1/F2 boundary that defines treatment "
  "eligibility.")

doc.add_heading("6. Limitations", level=1)
p("First, the external validation cohorts were modest in size (Fujiwara n = 213; UCAM "
  "n = 58), and UCAM was underpowered for F2 vs F3 (n = 12/7). Second, deconvolution "
  "relies on bulk reference atlases. Third, 149 of 649 signature genes (23.0%) were not "
  "measurable in the shared validation universe. Fourth, the TF regulon analysis was "
  "performed on the Fujiwara matrix and its single-cell localisation is descriptive. "
  "Fifth, the ferroptosis null findings apply to stage-defining bulk programmes and do "
  "not exclude pathogenic roles. Sixth, clinical covariates were limited to age, BMI and "
  "diabetes status. Seventh, paediatric samples were included in discovery and the sex "
  "interaction was inseparable from the paediatric cohort; age and sex were additionally "
  "unavailable for the largest discovery cohort. Eighth, per-stage-pair contrasts used "
  "limma-trend on the locked log2-CPM matrix rather than voom precision weights (stated "
  "deviation); the F3 vs F4 pair is low-powered (n = 73/16). Ninth, atlas cell types were "
  "assigned by canonical markers rather than the original publication's annotations. "
  "Tenth, the driver/suppressor balance was quantified as a score difference because a "
  "ratio is undefined where the denominator crosses zero.")

p("Conflict of Interest", bold=True)
p("The authors declare no conflicts of interest.")
p("Funding", bold=True)
p("This work was supported by [Funding sources to be added].")
p("Data Availability", bold=True)
p("All discovery and validation datasets are publicly available from GEO (GSE135251, "
  "GSE130970, GSE185051; Fujiwara SuperSeries GSE193084 with sub-series GSE192959, "
  "GSE193066, GSE193080) and the Kamzolas et al. (2026) data release; the single-cell "
  "atlas is GSE136103. Analysis code and processed data are available at [repository URL].")

doc.add_heading("References", level=1)
refs_start = find("Angulo, P., Kleiner")
for r in paras[refs_start:]:
    if r.startswith("Supplementary"):
        break
    par = p(r, size=10)
    par.paragraph_format.space_after = Pt(3)
# add the atlas reference
p("Ramachandran, P., Dobie, R., Wilson-Kanamori, J.R., et al., 2020. Resolving the "
  "fibrotic niche of human liver cirrhosis using single-cell transcriptomics. Nature 582, "
  "128–133.", size=10)

doc.add_heading("Supplementary Materials", level=1)
p("Supplementary tables: S1. Full differential-expression table (24,845 gene rows). S2. "
  "Complete enrichment ranking. S3. Benchmark: all panels × cohorts × three splits with "
  "coverage, AUROC, CI and DeLong comparisons. S4. Parsimony curve (13 sizes × 6 strata). "
  "S5. Longitudinal per-patient scores. S6. Clinical-subset composition. S7. Network "
  "communities and hubs. S8. NMF cluster-by-stage distributions. S9. Prioritisation. S10. "
  "Fixed operating-point performance. S11. Curated drug–gene edges. S12. Panel gene lists. "
  "S13. Cell-death pathway scoring and association tests (QuickGO sets). S14. "
  "Deconvolution: full correlation table. S15. SPIB post-hoc pairwise comparisons. S16. "
  "SPIB cohort-stratified analysis. S17. Top 20 prioritised transcription factors. S18. "
  "Per-stage-pair DGE tables and GSEA results. S19. Paired-biopsy ferroptosis scores and "
  "trajectory tests. S20. Co-expression network edges, null distribution and modules. "
  "S21. Effector-gene donor-pseudobulk values. S22. Load manifests and provenance records "
  "for all new analyses.", size=10)
p("Supplementary figures: Fig. S1. Cohort flow diagram. Fig. S2. Deconvolution results. "
  "Fig. S3. Sensitivity analysis. Fig. S4. Per-gene AUROCs. Fig. S5. SPIB "
  "cohort-stratified analysis. Fig. S7. Per-gene AUROCs across cohorts. Fig. S8. "
  "Extra-hepatic pathway enrichment. Fig. S9. Multi-omic fusion and SHAP interpretability. "
  "Fig. S10. Continuous pseudotime trajectory and longitudinal regression. Fig. S11. "
  "Ferroptosis co-expression network (permutation-nulled).", size=10)

doc.save(OUT)
print("saved", OUT)
