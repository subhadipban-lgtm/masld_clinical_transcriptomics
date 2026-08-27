#!/usr/bin/env python3
"""WS31 text pass — builds REVISED_v2 of the manuscript.

Base: the WS29-state REVISED builder lineage (paragraph sources from the original
2026-08-24 draft; tables from authoritative CSVs; figures from original media +
WS26-30 result figures). This pass applies, with numbers from files only:

  WS29: concurrency verdict (temporal framing removed); rank-25 claim replaced by
        uniform enrichment; rho-with-CIs cell-death framing; ZBTB2 demoted; set chain.
  WS30: Fujiwara matrisome-arm replication; gene-disjointness; compositional (not
        dilution) comparator folds; basement-membrane leading category; compact-panel
        compositional failure; twelve nulls; KG/GNN removals (Methods 2.7.6, S7, S11).
  WS30b-d: Fujiwara ferroptosis inversion = genome-wide mean-z drift artifact;
        discovery suppressor specificity marginal (empirical p 0.039) and UCAM
        at-threshold (0.04995), stated separately; drivers null-typical everywhere
        ("suppressor-biased" framing removed); drift-adjusted Table 7 (excess primary);
        random-set null described once as a general guard; degenerate genes-detected
        disclosure; no apoptosis-survives sentence; marginal results not combined.
"""
import html
import os

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SRC_TXT = "/tmp/manuscript_temporal.txt"
MEDIA = "/tmp/msx/word/media"
OUT = ("manuscript_assets/"
       "MASLD_Matrisome_Convergence_Threshold_2026-08-25_REVISED_v2.docx")

paras = [html.unescape(l) for l in open(SRC_TXT).read().split("\n")]
paras = [p.strip() for p in paras if p.strip()]
paras2 = [html.unescape(l) for l in open("/tmp/ms_rev.txt").read().split("\n")]
paras2 = [p.strip() for p in paras2 if p.strip()]

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
style.paragraph_format.line_spacing = 1.3
style.paragraph_format.space_after = Pt(6)
for s in ["Heading 1", "Heading 2", "Heading 3"]:
    st = doc.styles[s]
    st.font.name = "Times New Roman"
    st.font.color.rgb = RGBColor(0, 0, 0)


def p(text, bold=False, italic=False, size=None, align=None):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if align:
        par.alignment = align
    return par


def caption(text):
    par = p(text, italic=True, size=9)
    par.paragraph_format.space_after = Pt(12)
    return par


def figure(path, cap, width=6.1):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(path, width=Inches(width))
    caption(cap)


def table(df, cap, numfmt=None):
    caption(cap)
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, c in enumerate(df.columns):
        r = t.rows[0].cells[j].paragraphs[0].add_run(str(c))
        r.bold = True
        r.font.size = Pt(9)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, v in enumerate(row):
            txt = numfmt(v) if numfmt else str(v)
            r = cells[j].paragraphs[0].add_run(str(txt))
            r.font.size = Pt(9)
    for rowx in t.rows:
        for c in rowx.cells:
            for pp in c.paragraphs:
                pp.paragraph_format.space_after = Pt(2)
                pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def find(prefix, start=0):
    for i in range(start, len(paras)):
        if paras[i].startswith(prefix):
            return i
    raise ValueError("not found: " + prefix[:60])


def find2(prefix):
    for i, t in enumerate(paras2):
        if t.startswith(prefix):
            return t
    raise ValueError("not found in v1: " + prefix[:60])


# ============================ front matter ============================
TITLE = ("Independent hepatic transcriptional signatures converge on the matrisome and "
         "fail at the F2 treatment boundary in MASLD fibrosis: a multi-cohort "
         "benchmarking study")
p(TITLE, bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
for pref in ["[Author Name]", "[Affiliation]", "Corresponding author", "Email:"]:
    for j in [x for x in paras if x.startswith(pref)]:
        p(j, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

p("Abstract", bold=True, size=12)
p("Background & Aims: Fibrosis stage is the strongest histological determinant of "
  "liver-related outcomes in metabolic dysfunction-associated steatotic liver disease "
  "(MASLD). Multiple hepatic transcriptional signatures of stage have been published, but "
  "they are rarely validated in fully independent cohorts, almost never compared "
  "head-to-head under a common protocol, and have not been evaluated at the fibrosis "
  "thresholds that now govern treatment eligibility. We benchmarked our signature against "
  "five published panels under a locked protocol, characterised the biological programme "
  "they share, and asked whether the ferroptosis programme — frequently invoked in MASLD "
  "fibrogenesis — tracks fibrosis stage once guarded against genome-wide expression drift.")
p("Methods: Differential expression (voom-limma) between late (F3–F4, n = 89) and early "
  "(F0–F2, n = 260) fibrosis across 349 biopsy-staged patients from three cohorts; "
  "validation in 271 patients across two independent cohorts (Fujiwara n = 213, UCAM "
  "n = 58) under a locked normalisation; benchmarking at F0–2 vs F3–4, F0–1 vs F2–4 and "
  "F2 vs F3. Matrisome (MSigDB v2023.2.Hs) and FerrDb ferroptosis programmes were "
  "evaluated cross-sectionally and per adjacent stage transition (limma-trend contrasts), "
  "with replication in Fujiwara. Every mean-z pathway score was guarded by a random-set "
  "null (1,000 size-matched sets, identical estimator) and reported as a drift-adjusted "
  "excess. Effector-gene localisation used a 62,210-cell single-cell atlas across 10 "
  "donors with donor-level pseudobulk inference.")
p("Results: The signature (649 genes, 500 measurable) discriminated late from early "
  "fibrosis at AUROC 0.811 (Fujiwara) and 0.826 (UCAM); the 75-gene panel reached 0.847 "
  "and 0.769, above all five published comparators in Fujiwara. All five comparator panels "
  "and our signature were enriched for matrisome genes (fold 6.3–16.2); the higher "
  "comparator folds are compositional, not dilution — each sits at the 90th–100th "
  "percentile of size-matched subsamples of our signature. Stage-resolved enrichment "
  "shows the matrisome programme present from the earliest transition (F0 vs F1) and "
  "rising monotonically; the rising trend replicated in Fujiwara. The two programmes are "
  "gene-disjoint (zero leading-edge overlap). Ferroptosis, by contrast, is a "
  "well-characterised negative: the suppressor-set stage association is specific in "
  "discovery only marginally (random-set empirical p = 0.039), exactly at threshold in "
  "UCAM (p = 0.04995, reported separately), and an artifact of genome-wide mean-z drift "
  "in Fujiwara (33rd percentile of the size-matched null; absorbed by depth covariates); "
  "the driver set is null-typical in both cohorts; a ferroptosis-DEG panel fails the "
  "treatment boundary (AUROC 0.588, CI including 0.50); and no pro-ferroptotic effector "
  "gene localises to stellate cells across donors. Twelve rigorously evaluated nulls are "
  "reported.")
p("Conclusions: Independently derived MASLD fibrosis signatures converge on a shared, "
  "matrisome-dominated programme that is present from the earliest stage transition and "
  "strengthens with fibrosis — and all of them fail at the F2 treatment boundary, a "
  "biological ceiling of bulk transcriptomics not rescued by model class, covariate "
  "fusion or panel compression. The ferroptosis programme does not carry staging "
  "information beyond genome-wide drift. Two methodological guards — compositional "
  "correction for deconvolution and a random-set null for pathway scoring — each changed "
  "a conclusion in our own data and are offered as practice. Clinical deployment is not "
  "advocated on present evidence.")
p("Keywords: MASLD; fibrosis staging; transcriptional signature; matrisome; treatment "
  "eligibility; benchmarking; ferroptosis; random-set null", italic=True)

# ============================ introduction ============================
doc.add_heading("1. Introduction", level=1)
p(paras[find("Metabolic dysfunction-associated steatotic liver disease (MASLD) is now")])
p("The transition from F0 to F1 represents the earliest detectable molecular shift in "
  "MASLD fibrosis, yet its drivers remain incompletely characterized. Ferroptosis — an "
  "iron-dependent, non-apoptotic cell death programme — has been implicated in early "
  "metabolic injury (Peleman et al., 2024a, 2024b; Zhou et al., 2026), and inhibiting "
  "ferroptosis prevents progression of steatotic liver disease in obese mice (Park et "
  "al., 2024). Whether the ferroptosis transcriptional programme tracks fibrosis stage "
  "in human bulk tissue — as opposed to riding a genome-wide expression drift, to which "
  "any mean-z gene-set score is susceptible — is the question this paper answers "
  "directly, with the null guard that question requires.")
for pref in ["At later stages, fibrosis progression", "Integrative multi-omics studies",
             "While multiple hepatic transcriptional signatures"]:
    p(paras[find(pref)])
p("We addressed these gaps directly. We derived a fibrosis-stage signature from 349 "
  "biopsy-staged patients under a locked pipeline; validated it in 271 patients across "
  "two verified-independent cohorts; benchmarked it against five published panels at "
  "three clinically defined thresholds including the treatment-eligibility boundary; "
  "characterised the shared matrisome programme cross-sectionally, per stage transition "
  "and against a replication cohort; and evaluated the ferroptosis programme under a "
  "random-set null guard at every step. The analysis protocol was pre-registered and all "
  "validation matrices were locked before external cohort evaluation.")

# ============================ methods ============================
doc.add_heading("2. Methods", level=1)
doc.add_heading("2.1 Cohorts", level=2)
p("Three bulk RNA-seq cohorts with verified histological staging were harmonised into "
  "the discovery set: GSE135251 (n = 216), GSE130970 (n = 78) and GSE185051 (paediatric, "
  "n = 55; 349 patients total). Two additional GEO series (GSE167523, GSE126848) were "
  "excluded because their binary NASH labels could not be mapped to histological stage. "
  "External validation used the Fujiwara cohort (n = 213; SuperSeries GSE193084, "
  "comprising GSE192959, GSE193066 and GSE193080) and the UCAM cohort (n = 58). All "
  "cohorts were biopsy-staged (NASH CRN, F0–F4) and verified to share no samples with "
  "the discovery set (Fig. S1).")
doc.add_heading("2.2 Processing", level=2)
p(paras[find("Raw counts were filtered")])
doc.add_heading("2.3 Differential expression", level=2)
p(paras[find("voom precision weights + lmFit")])
doc.add_heading("2.4 Signature scoring and thresholds", level=2)
p(paras[find("Each gene was z-scored")])
doc.add_heading("2.5 Comparator panels", level=2)
p(find2("Five published fibrosis-stage panels from the Kamzolas"))
doc.add_heading("2.6 Panel-size selection", level=2)
p(paras[find("The pre-specified rule was")])
doc.add_heading("2.7 Exploratory data modelling", level=2)
for h, pref in [("2.7.1 Pseudotime Trajectory Mapping", "PCA was fitted on baseline"),
                ("2.7.2 Matrisome enrichment analysis", "Hypergeometric tests evaluated"),
                ("2.7.3 Ferroptosis enrichment", "Gene-set enrichment of the signed-t"),
                ("2.7.4 Transcription factor regulon analysis", "TFPLACEHOLDER"),
                ("2.7.5 TRRUST network analysis", "TF-target interactions were filtered to those")]:
    doc.add_heading(h, level=3)
    if "PLACEHOLDER" in pref:
        p("Transcription factor (TF) regulon inference was performed using pySCENIC "
          "(Aibar et al., 2017) with AUCell scoring against the CollecTRI regulatory "
          "network on the Fujiwara validation matrix. The top 20 TFs by activity "
          "variance were prioritized. Differential TF activity was tested with "
          "Mann–Whitney U (Stage 3, n = 30, vs Stage 0, n = 12) with Benjamini–Hochberg "
          "correction across TFs; post-hoc pairwise tests for SPIB were stratified by "
          "cohort. Single-cell localisation used the atlas of Section 2.7.10.")
    else:
        p(paras[find(pref)])
p("MSigDB release pinning: all matrisome analyses use MSigDB C2 v2023.2.Hs "
  "(NABA_MATRISOME: 1,026 genes, 627 in the tested background). A local NABA file of "
  "different provenance used in an early round intersects the same 627 background genes "
  "and produced no value reported here.")
doc.add_heading("2.7.6 Ferroptosis programme across the continuum", level=3)
p("Ferroptosis drivers (FerrDb V2 filtered, 264) and suppressors (238) were scored in "
  "every cohort as the mean within-sample z-score of set members present in the locked "
  "matrix. The gene-set size chain is stated once: raw FerrDb V2 exports 2,146/3,047 → "
  "filtered 264/238 → present in the locked 12,537-gene matrix 194/186 → present in the "
  "15,223-gene tested background 201/184 → hit in the stage-pair ranked lists 185/178 "
  "(each count is cited where used). A ferroptosis-DEG classifier (9 genes) was scored "
  "with discovery log2FC signs; AUROCs carry class-stratified bootstrap 95% CIs and the "
  "discovery row is labelled in-sample. Age and sex associations were tested within "
  "cohorts.")
doc.add_heading("2.7.7 Random-set null guard (applies to every mean-z score)", level=3)
p("Because a mean-z gene-set score inherits any genome-wide drift between expression and "
  "stage, every pathway score in this paper is tested against 1,000 random gene sets of "
  "the same size drawn from the same matrix and scored with the identical estimator "
  "(seed 42). We report the empirical two-sided p and the drift-adjusted excess — the "
  "observed Spearman ρ minus the null centre — as the primary effect size, because null "
  "centres scale with set size (a 1,190-gene random set is essentially the drift "
  "itself). A score is called specific only when its empirical p < 0.05.")
doc.add_heading("2.7.8 Stage-resolved enrichment, paired biopsies, and technical covariates", level=3)
p("For each adjacent stage pair in each cohort, limma with eBayes(trend = TRUE) was "
  "fitted on the locked log2-CPM matrix (limma-trend; the pooled discovery contrast "
  "additionally used voom precision weights on raw counts, a stated deviation). "
  "Weighted running-sum GSEA (1,000 permutations, seed 42) tested the ten NABA "
  "categories and two FerrDb sets against each pair's t-ranked list, with one BH family "
  "across 12 sets × 4 transitions per cohort. Paired second biopsies (58 patients, "
  "timepoint order verified against locked baseline scores, ρ = 0.910 vs 0.668) were "
  "scored identically; trajectory groups follow Δfibrosis. Technical covariates: the "
  "locked Fujiwara matrix is post-normalisation GEO log2 expression with no zero "
  "entries, so genes-detected is degenerate (12,537 in every sample) and true library "
  "size is unrecoverable for the 213 cross-sectional samples; we therefore adjusted for "
  "recoverable proxies (all-gene mean-z; linear-space expression sum; mean log2 "
  "expression) and, for real raw-count library size and genes detected, used the n = 40 "
  "baseline subset matched to the locked build.")
doc.add_heading("2.7.9 Cell-death context and co-expression network", level=3)
p(find2("Apoptosis (GO:0006915)") + " All pathway scores are reported under the "
  "random-set null guard of Section 2.7.7 (drift-adjusted excess as the primary "
  "quantity).")
doc.add_heading("2.7.10 Single-cell effector-gene localisation", level=3)
p(find2("The GSE136103 atlas"))
doc.add_heading("2.8 Sensitivity Analysis and Signature Stability", level=2)
p(paras[find("DEG identification was re-run")])
doc.add_heading("2.9 Bulk deconvolution", level=2)
p(paras[find("Cell-type associations were assessed by non-negative")])
doc.add_heading("2.10 Statistics", level=2)
p(paras[find("AUROCs with 95% confidence intervals")])
p("Bootstrap 95% confidence intervals (2,000 class-stratified patient-level resamples, "
  "seed 42) accompany every AUROC introduced in Sections 2.7.6–2.7.10. Stochastic "
  "analyses report the seed and permutation count used.")

# ============================ results ============================
doc.add_heading("3. Results", level=1)
doc.add_heading("3.1 Discovery of the 649-gene fibrosis signature", level=2)
p(paras[find("Differential expression between late (F3–F4, n = 89)")])
t1 = pd.DataFrame({"Characteristic": ["Total patients (discovery)", "Late fibrosis (F3–F4)",
                                      "Early fibrosis (F0–F2)", "Validation: Fujiwara (GSE193084)",
                                      "Validation: UCAM", "Genes tested (voom)",
                                      "Signature genes identified",
                                      "Measurable in validation universe"],
                   "Value": [349, 89, 260, 213, 58, "15,223", 649, "500 (77.0%)"]})
table(t1, "Table 1. Discovery cohort characteristics.")
figure(f"{MEDIA}/994f01431ec1dd207157fc02d578b510e9683ebf.png",
       "Figure 1. Discovery differential expression and matrisome enrichment. (A) Volcano "
       "plot of Late (F3–F4, n = 89) versus Early (F0–F2, n = 260). The signature is "
       "uniformly matrisome-enriched across its ranking: the top ten genes carry "
       "matrisome membership at 4.86-fold over background — marginally above the "
       "signature-wide 4.68-fold — with enrichment peaking at 7.3–7.8-fold across ranks "
       "25–100 (matrisome_by_rank_fold.csv). (B) NABA matrisome enrichment fold across "
       "the ten canonical categories (NABA_MATRISOME fold 4.68, padj = 1.74 × 10⁻⁴⁹).")

doc.add_heading("3.2 Transcription factor regulon analysis reveals stage-specific regulatory shifts", level=2)
p(paras[find("To identify upstream regulators of the fibrotic program")])
p(find2("Stage-dependent TF activity. Mann"))
d3 = pd.read_csv("results/ws19/diff_activity_S3_vs_S0_sig.csv")
d3 = d3.rename(columns={"Mean_Stage0": "Stage-0 mean", "Mean_Stage3": "Stage-3 mean",
                        "Mean_Diff (S3-S0)": "Δ (S3−S0)", "P_value": "p (MWU)",
                        "FDR_adj_P": "FDR"})
table(d3[["TF", "Stage-0 mean", "Stage-3 mean", "Δ (S3−S0)", "p (MWU)", "FDR"]],
      "Table 3. Differentially active transcription factors (Mann–Whitney U, Stage 3 vs "
      "Stage 0, n = 30/12; BH across tested regulons). Source: diff_activity_S3_vs_S0_sig.csv.",
      numfmt=lambda v: f"{v:.4g}" if isinstance(v, float) else v)
p(paras[find("SPIB shows consistent stage-dependent activity")])
figure(f"{MEDIA}/9fe538c56a0f6617b1bb582abda7db89d731e0ce.png",
       "Figure 3. TF regulon analysis: patient clustermap of top 20 transcription factors.")
figure(f"{MEDIA}/371e931d1b0c3cc4a45425c5f9a2814d88d200c7.png",
       "Figure 3C. SPIB activity by fibrosis stage and cohort. Left: GSE193066 (n = 106, "
       "Kruskal-Wallis p = 1.15 × 10⁻⁵). Right: GSE192959 (n = 42, p = 0.288). SPIB "
       "differentiates all stage pairs including F0 vs F1 (FDR = 0.018).")

doc.add_heading("3.3 A TRRUST-documented ZBTB2–TP53 axis (hypothesis-generating)", level=2)
p("TRRUST v2 documents ten interactions for the 20 prioritised TFs (Table 4); all four "
  "ZBTB2 edges derive from a single publication (PMID 19380588), and only three of its "
  "four targets are FerrDb ferroptosis genes (CDKN1A, MDM2, TP53 — the p53 axis, which "
  "is annotated across several stress-response gene sets). The enrichment of FerrDb "
  "genes among ZBTB2 targets is nominally significant (fold 10.62, p = 0.0013, FDR "
  "0.0066), but the same targets are equally enriched in FerrDb drivers and suppressors, "
  "nominally enriched for apoptosis and necroptosis, and absent from GO:0097707 "
  "ferroptosis (0 of 4; WS30 zbtb2_specificity.csv). With n = 4 targets and a single "
  "supporting PMID, this is a hypothesis-generating observation of a ZBTB2–TP53 axis, "
  "not evidence that ZBTB2 regulates ferroptosis.")
t4 = pd.read_csv("/tmp/trrust_table4_full.csv").rename(columns={"ferro_target": "Ferroptosis target"})
t4["Ferroptosis target"] = t4["Ferroptosis target"].map({True: "Yes", False: "No"})
table(t4, "Table 4. TRRUST-documented interactions of the 20 prioritised TFs (all edges "
          "with recorded Mode and PMID). Only ZBTB2's edges target FerrDb genes.")
figure(f"{MEDIA}/924e02737fc452cd5004b95531f650f5b7994569.png",
       "Figure 4. TRRUST-documented ZBTB2 edges (hypothesis-generating; see text).")

doc.add_heading("3.4 Signature convergence with an independently derived panel", level=2)
p(find2("Testing the 649-gene signature for membership"))

doc.add_heading("3.5 External validation and panel benchmarking", level=2)
p(find2("The 649-gene signature (500 measurable genes) was validated"))
t3f = pd.read_csv("results/ws15/table3_final.csv")
piv = t3f[~t3f.apparent].pivot_table(index=["split", "cohort"], columns="panel", values="AUROC").reset_index()
piv.columns.name = None
order = ["Ours_full", "Ours_75", "Ours_10", "Kamzolas_145", "Kamzolas_57BM",
         "Kamzolas_15BM", "Kamzolas_194PT", "Kamzolas_3gene"]
piv = piv[["split", "cohort"] + order].rename(
    columns={"Ours_full": "Ours-649", "Ours_75": "Ours-75", "Ours_10": "Ours-10",
             "Kamzolas_145": "Kamz-145", "Kamzolas_57BM": "Kamz-57BM",
             "Kamzolas_15BM": "Kamz-15", "Kamzolas_194PT": "Kamz-194", "Kamzolas_3gene": "Kamz-3"})
table(piv, "Table 2. External validation AUROCs (held-out strata; rebuilt from "
          "table3_final.csv, WS15 locked build).", numfmt=lambda v: f"{v:.3f}" if isinstance(v, float) else v)
figure(f"{MEDIA}/1ea1e928892ed0fae6c235e23e77937325a30646.png",
       "Figure 2. External validation and panel benchmarking at three clinically defined thresholds.")

doc.add_heading("3.6 Performance at clinically defined thresholds", level=2)
p(find2("Performance degraded for every panel"))
figure(f"{MEDIA}/db9f72f8eac877da53e74ad0a4f448444c1e88af.png",
       "Figure 6. Prioritisation and panel characterisation at treatment thresholds.")
doc.add_heading("3.6.1 Multi-omic fusion performance", level=2)
p(find2("Early fusion of transcriptomic signatures"))
doc.add_heading("3.7 Panel size is not identifiable from performance data", level=2)
p(paras[find("Ranking the measurable signature by discovery |t|")])
p(paras[find("The decisive observation is not the ordering")])
p(paras[find("We therefore report that panel size is not identifiable")])
figure(f"{MEDIA}/c17b01ca276aa4a65f25990056aadd36669c413d.png",
       "Figure 5. Panel size parsimony analysis.")
doc.add_heading("3.8 Fixed threshold is only partially transportable", level=2)
p(paras[find("A fixed operating point derived")])
doc.add_heading("3.9 Comparison with clinical scores: parity with FIB-4", level=2)
p(paras[find("In a 48-patient clinical subset")])
t5 = pd.DataFrame({"Score": ["FIB-4", "APRI", "Ours (5-gene)"],
                   "AUROC": [0.555, 0.592, 0.561],
                   "95% CI": ["0.366–0.732", "0.416–0.756", "0.375–0.733"], "n": [47, 47, 47]})
table(t5, "Table 5. Clinical score comparison (48-patient matched subset; DeLong p > 0.35).")
figure(f"{MEDIA}/339a7d88ce3b893e796a3262d529385b7d8dced4.png",
       "Figure 7. Benchmarking against clinical scores in the matched subset.")
doc.add_heading("3.10 Longitudinal analysis: technical artefact", level=2)
p(find2("In 58 paired biopsies, the signature score did not track"))
figure(f"{MEDIA}/6f45f425a57b24b9b8254e23e1bb16cb23137eaa.png",
       "Figure 8. Longitudinal pseudotime analysis.")
doc.add_heading("3.11 Cell-type localisation: compositional correction required", level=2)
p(paras[find("Deconvolution revealed a cholangiocyte-dominant association")])
doc.add_heading("3.12 Exploratory analyses", level=2)
doc.add_heading("3.12.1 WGCNA", level=3)
p(paras[find("A properly re-run weighted co-expression analysis")])
doc.add_heading("3.12.2 Ferroptosis", level=3)
p(paras[find("Gene-set enrichment of the signed-t ranking")])
doc.add_heading("3.12.3 Sex", level=3)
p(paras[find("Clean sex labels were available")])
doc.add_heading("3.12.4 Extra-hepatic pathways", level=3)
p(paras[find("Only REACTOME_HEMOSTASIS was enriched")])

# ---- 3.13 ferroptosis, null-guarded ----
doc.add_heading("3.13 The ferroptosis programme under its random-set null", level=2)
p("In discovery, the suppressor-set score correlates with stage at ρ = 0.286 — but the "
  "parametric Kruskal–Wallis p (1.2 × 10⁻⁶) tests the wrong null: it asks whether the "
  "score varies with stage, not whether ferroptosis genes do so more than 178 random "
  "genes. Against 1,000 size-matched random sets the observed ρ sits at the 97.3rd "
  "percentile of a null centred at +0.042 (empirical two-sided p = 0.039): specific, "
  "and marginal. In UCAM the same test gives an empirical p of exactly 0.04995 "
  "(95.1st percentile of a null centred at +0.264) — at the pre-committed threshold, "
  "not beyond it. We state the two results separately and draw no combined claim: "
  "neither is a strong result. In Fujiwara the raw association is negative (ρ = −0.224) "
  "but sits at the 33rd percentile of a size-matched null centred at −0.128 (empirical "
  "p = 0.35): random gene sets behave like the 'ferroptosis' result, which is therefore "
  "a genome-wide mean-z drift, not a ferroptosis signal. Consistent with that reading, "
  "the association attenuates to −0.075 (p = 0.28) under a library-depth proxy that the "
  "score tracks at ρ = 0.72, and to −0.300/-0.265 (borderline) under real raw-count "
  "library size and genes-detected in the n = 40 subset. The driver set is null-typical "
  "in both cohorts (57th percentile in discovery, p = 0.59; 65th in UCAM, p = 0.36): "
  "there is no supported driver–suppressor asymmetry, and we make no claim of one. A "
  "nine-gene ferroptosis-DEG classifier reached 0.659 (F3+) in Fujiwara but failed the "
  "treatment boundary at 0.588 (CI 0.499–0.673, including 0.50; Table 6). Age "
  "associations were cohort-confounded and null within the adult cohort; sex showed no "
  "association where measurable.")
figure("results/ws26_ferroptosis_additional/fig_2_2_gsva_boxplots.png",
       "Figure 9. Ferroptosis driver and suppressor mean-z scores across fibrosis stages, "
       "discovery 349. Interpretation requires the random-set null (Section 3.13): only "
       "the suppressor arm is specific, marginally.")
t6 = pd.read_csv("results/ws26_ferroptosis_additional/ferroptosis_signature_auroc.csv")
t6 = t6.rename(columns={"CI95_lo": "CI lo", "CI95_hi": "CI hi", "n_positive": "n pos",
                        "genes_measured": "genes", "in_sample_discovery": "in-sample"})
t6["in-sample"] = t6["in-sample"].map({True: "Yes", False: "No"})
table(t6[["cohort", "threshold", "AUROC", "CI lo", "CI hi", "n", "n pos", "genes", "in-sample"]],
      "Table 6. Ferroptosis-DEG classifier AUROCs with bootstrap 95% CIs.",
      numfmt=lambda v: f"{v:.3f}" if isinstance(v, float) else v)

# ---- 3.14 stage-resolved: concurrency ----
doc.add_heading("3.14 Stage-resolved enrichment: the matrisome programme is early, rising, and replicated; ferroptosis is discovery-only", level=2)
p("Per-adjacent-stage-pair GSEA (one BH family of 48 per cohort) shows the two "
  "programmes are concurrent, not sequential. Matrisome categories are already "
  "significant at the earliest transition (F0 vs F1: NABA_PROTEOGLYCANS padj 0.021, "
  "NABA_ECM_REGULATORS 0.028) and rise monotonically — NABA_MATRISOME NES 1.40, 1.16, "
  "2.29, 2.35 across the four transitions — while the ferroptosis suppressor set peaks "
  "at F0 vs F1 (NES 1.62) and is absent thereafter. In Fujiwara the matrisome arm "
  "replicates as a rising trend (NES 0.83, −1.47, 1.85, 1.60; Fig. 11), with F0 vs F1 "
  "low-powered (n = 12/58); the ferroptosis arm does not replicate, and Section 3.13 "
  "shows its apparent Fujiwara inversion is drift, not biology. The pre-committed "
  "concurrency rule therefore fired: no temporal ordering between the programmes is "
  "claimed. The programmes are additionally gene-disjoint: the suppressor set's F0vsF1 "
  "leading edge (62 genes) shares zero genes with every significant matrisome F2vsF3 "
  "leading edge, and only 5/201 drivers and 3/184 suppressors are annotated to any NABA "
  "category.")
figure("results/ws27_ferroptosis_followups/fig_stagepair_gsea.png",
       "Figure 10. Ferroptosis set enrichment per adjacent stage transition (discovery; "
       "family of 48 including the matrisome sets of Figure 11; * padj < 0.05).")
figure("results/ws30/stagepair_discovery_vs_fujiwara.png",
       "Figure 11. Stage-resolved enrichment, discovery vs Fujiwara (ringed = padj < 0.05 "
       "in the family of 48; LP = low-powered F0 vs F1, n = 12/58). The matrisome rising "
       "trend replicates; the ferroptosis arm does not.")
doc.add_heading("3.15 Paired biopsies: ferroptosis scores do not track fibrosis trajectory", level=2)
p(find2("In the 58 patients with paired biopsies (timepoint order verified"))
figure("results/ws27_ferroptosis_followups/fig_paired_trajectory.png",
       "Figure 12. Ferroptosis score change (follow-up − baseline) by fibrosis trajectory, n = 58.")

# ---- 3.16 cell-death, drift-adjusted ----
doc.add_heading("3.16 Cell-death programmes and network structure under the null guard", level=2)
p("Table 7 reports the seven cell-death pathways with the drift-adjusted excess "
  "(ρ_obs − null centre) as the primary column, because raw ρ mixes signal with a "
  "size-dependent drift offset (null centres range from +0.004 at 3 genes to +0.112 at "
  "1,190). On that scale necroptosis carries the largest excess (+0.349, empirical p = "
  "0.006), followed by GO ferroptosis (+0.333, p = 0.017) and the FerrDb suppressor set "
  "(+0.241, p = 0.039). Autophagy is a borderline miss (excess +0.180, p = 0.078) and is "
  "reported as such; pyroptosis (3 genes, p = 0.26) and the FerrDb driver set (p = 0.59) "
  "are null-typical and are demoted to the supplementary table. The bootstrap CI on the "
  "necroptosis-minus-ferroptosis ρ difference spans zero ([−0.052, 0.100]): no cell-death "
  "programme is separable from another, and none dominates. The co-expression network of "
  "the 347 measurable ferroptosis genes is massive but weakly structured (14,917 edges "
  "at |ρ| > 0.5 against a per-gene permutation null of zero; modularity Q = 0.13), with "
  "driver–suppressor cross-wiring in excess of random wiring (0.55 vs 0.46, p ≈ 10⁻¹⁶⁶) "
  "— coordinated regulation, not opposing arms.")
nul = pd.read_csv("results/ws30/random_set_null_pathways.csv")
d = nul[nul.cohort == "Discovery"].copy()
d["excess"] = d.rho_obs - d.null_mean
d = d.sort_values("excess", ascending=False)
d["status"] = d.verdict.str.replace(" (drift)", "", regex=False)
t7 = d[["pathway", "n_genes", "rho_obs", "null_mean", "excess",
        "obs_percentile", "empirical_p_two_sided", "status"]].rename(
    columns={"pathway": "pathway", "n_genes": "genes", "rho_obs": "ρ obs",
             "null_mean": "null centre", "excess": "excess (primary)",
             "obs_percentile": "pct of null", "empirical_p_two_sided": "emp. p",
             "status": "verdict"})
table(t7, "Table 7. Cell-death pathway stage association under the random-set null "
          "(discovery 349; primary column = drift-adjusted excess). Pyroptosis and FerrDb "
          "drivers are null-typical (supplementary). Source: random_set_null_pathways.csv.",
      numfmt=lambda v: f"{v:.3f}" if isinstance(v, float) else v)
figure("results/ws28_celldeath/fig_celldeath_comparison.png",
       "Figure 13. Cell-death pathway stage association, unadjusted view (see Table 7 for "
       "the drift-adjusted primary quantities).")

# ---- 3.17 effectors ----
doc.add_heading("3.17 Effector-gene localisation: no stellate-cell enrichment", level=2)
p(find2("In the GSE136103 atlas (62,210 cells"))
figure("results/ws28_sc_effectors/fig_effector_celltype_heatmap.png",
       "Figure 14. Ferroptosis effector genes by cell type (GSE136103; descriptive "
       "cell-level means — inferential tests use donor pseudobulk, Table 8).")
t8 = pd.read_csv("results/ws28_sc_effectors/effector_HSC_vs_hepatocyte.csv").rename(
    columns={"log2_HSC_over_hep": "log2 HSC/hep", "wilcoxon_p": "p", "p_adj": "padj",
             "n_donors": "donors"})
table(t8[["gene", "HSC_mean", "hep_mean", "log2 HSC/hep", "p", "padj", "donors"]],
      "Table 8. Effector genes, stellate vs hepatocyte donor pseudobulk (n = 10 donors; "
      "BH across genes).", numfmt=lambda v: f"{v:.3g}" if isinstance(v, float) else v)

# ---- 3.18 matrisome characterisation ----
doc.add_heading("3.18 Characterising the convergence: composition, not dilution", level=2)
p("The higher matrisome folds of the comparator panels are compositional, not a size "
  "effect: subsampling our 649-gene signature to each comparator's size 1,000 times "
  "yields median folds of 4.75–4.86 at every size, and every comparator sits at the "
  "90.1st–100th percentile of its size-matched distribution (Kamzolas 145 at the "
  "100th). Within the signature, NABA_BASEMENT_MEMBRANES is the leading category (12 "
  "genes, fold 8.28, padj 1.3 × 10⁻⁸), of which only LAMC2 (rank 46) and LAMC3 (rank "
  "48) reach the top 100 by |t| and the 75-gene panel; none are cholangiocyte-associated "
  "in the atlas. The compact 10-gene panel's failure to enrich (2 of 10, padj 0.301) is "
  "compositional rather than a power artifact — significance at that size requires 3 of "
  "10, and the 3-gene comparator is significant with 2 of 3. The full 8 × 10 "
  "panel-by-category matrix is provided as a supplementary table (v2023.2.Hs throughout).")

# ============================ discussion ============================
doc.add_heading("4. Discussion", level=1)
doc.add_heading("Matrisome convergence is the finding", level=2)
p("Every independently derived signature we tested — five published panels and our own "
  "— is enriched for the same matrisome programme (fold 6.3–16.2), from the top of the "
  "discriminating-gene ranking to its base, present at the earliest stage transition and "
  "rising with fibrosis in both the discovery and the Fujiwara cohort. The comparator "
  "panels are not merely smaller versions of ours: they are genuinely more matrix-dense "
  "at equal size (90th–100th percentile of size-matched subsamples). The programmes are "
  "gene-disjoint from the ferroptosis signal. This convergence aligns with the "
  "mechanosensitive feed-forward loop of matrisome remodelling — LOX-mediated "
  "crosslinking, integrin mechanosensing, YAP/TAZ activation (Rudolph & Chin, 2024; Zhao "
  "et al., 2025; Fan et al., 2024; Benavides et al., 2026) — and with the basement-"
  "membrane category leading the enrichment.")
doc.add_heading("Ferroptosis is a well-characterised negative", level=2)
p("Guarded against genome-wide drift, the ferroptosis programme does not carry staging "
  "information. The suppressor-set association is discovery-specific and marginal "
  "(empirical p = 0.039), at-threshold in UCAM (0.04995) — two weak results stated "
  "separately, not combined — and in Fujiwara it is indistinguishable from drift (33rd "
  "percentile of the size-matched null), attenuating under depth covariates as this "
  "cohort's paired-biopsy analysis already warned. The driver arm is null-typical "
  "everywhere, so no driver–suppressor asymmetry is claimed. A ferroptosis-only panel "
  "fails the F2 boundary with a CI that includes chance; paired-biopsy ferroptosis "
  "scores do not track trajectory; and no pro-ferroptotic effector gene localises to "
  "stellate cells across donors, with iron-import machinery trending lower there. What "
  "survives is a hypothesis-generating ZBTB2–TP53 axis (n = 4 targets, one PMID) and a "
  "suppressor-set trend — observations, not findings. This matters beyond one paper: "
  "the ferroptosis-staging narrative is currently widespread (Peleman et al., 2024b; "
  "Wang, 2026), and our data suggest that in bulk human liver it does not survive a "
  "size-matched null.")
doc.add_heading("At the F2 treatment threshold, the transcriptional signal reaches a biological ceiling", level=2)
p("All signatures fail to resolve the F1→F2 transition. At F0–1 vs F2–4 the "
  "best-performing panel achieved an AUROC of 0.786, and at F2 vs F3 the evidence rests "
  "on a single cohort (0.746). Multiple independently derived signatures converge on the "
  "same matrisome programme and all fail at these boundaries, suggesting the molecular "
  "distinction between stages is not robustly captured by tissue-level gene expression. "
  "Alternative approaches — single-cell resolution, spatial transcriptomics, or "
  "protein-based biomarkers — may be required at the eligibility boundary (Bong & Seo, "
  "2026; Yang et al., 2026; Ko et al., 2026).")
p(paras[find("Discrimination was flat across panel sizes")])
p(paras[find("The deconvolution analysis revealed a cholangiocyte-dominant")])
doc.add_heading("Two methodological guards that changed conclusions in our own data", level=2)
p("The compositional correction for deconvolution-based cell-type claims removed an "
  "apparent stellate-cell association that was hepatocyte-complement artefact. The "
  "random-set null for mean-z pathway scoring caught a genome-wide drift that would "
  "otherwise have entered the literature as a ferroptosis finding — in opposite "
  "directions in two different cohorts (negative in Fujiwara, positive in UCAM, where "
  "null centres reach +0.30). Neither guard is standard practice; both are one-line "
  "additions to standard analyses; and each flipped a conclusion in our own hands. We "
  "suggest both as defaults for bulk-transcriptomic pathway scoring and deconvolution.")
doc.add_heading("Matrisome remodelling, mechanotransduction, and matrikines", level=2)
p(paras[find("The convergent matrisome programme identified here")])
doc.add_heading("Twelve rigorously evaluated null results", level=2)
p("The signature added no prognostic information beyond baseline stage in 58 paired "
  "biopsies. The fixed discovery-derived threshold transported only partially. Multi-omic "
  "fusion failed to improve treatment-boundary discrimination in all five architectures. "
  "A continuous pseudotime trajectory did not detect therapeutic regression, and its "
  "timepoint effect was a technical artefact after covariate adjustment. A re-run WGCNA "
  "found no module tracking stage. Ferroptosis genes were not enriched in the signature, "
  "did not discriminate at F2 as a standalone panel, did not track paired-biopsy "
  "trajectory, and did not localise to stellate cells as effectors. Graph centrality did "
  "not associate with GWAS membership (Mann–Whitney p = 1.0, n = 64), and ZBTB2's "
  "targets are not specifically ferroptosis genes. A sex×stage interaction did not "
  "survive restriction to adult cohorts.")

doc.add_heading("5. Conclusions", level=1)
p("Independently derived hepatic transcriptional signatures of MASLD fibrosis converge "
  "on a shared matrisome programme — present from the earliest stage transition, rising "
  "with fibrosis, replicated in an independent cohort, and compositional rather than a "
  "panel-size artifact — and every one of them fails at the F2 treatment boundary that "
  "governs eligibility. That failure is a biological ceiling of bulk transcriptomics: it "
  "is not rescued by model class, clinical covariate fusion, or panel compression, and "
  "there is no performance-based argument for any panel size.")
p("The ferroptosis programme, rigorously guarded, does not carry staging information "
  "beyond genome-wide expression drift; we report it as a characterised negative with "
  "two hypothesis-generating observations. Methodologically, we offer two guards — "
  "compositional correction for deconvolution and a size-matched random-set null for "
  "pathway scoring — each of which changed a conclusion in our own data. Clinical "
  "deployment of bulk transcriptomic staging is not advocated on present evidence; the "
  "field should look to single-cell, spatial, or protein-based measurement at the "
  "eligibility boundary.")

doc.add_heading("6. Limitations", level=1)
p("First, the external validation cohorts were modest (Fujiwara n = 213; UCAM n = 58), "
  "and UCAM was underpowered for F2 vs F3 (n = 12/7). Second, deconvolution relies on "
  "bulk reference atlases. Third, 149 of 649 signature genes (23.0%) were not measurable "
  "in the shared validation universe. Fourth, the TF regulon analysis was performed on "
  "the Fujiwara matrix and its single-cell localisation is descriptive. Fifth, the "
  "ferroptosis nulls apply to stage-defining bulk programmes and do not exclude "
  "pathogenic roles. Sixth, clinical covariates were limited to age, BMI and diabetes "
  "status. Seventh, paediatric samples were included in discovery and the sex "
  "interaction was inseparable from the paediatric cohort; age and sex were unavailable "
  "for the largest discovery cohort. Eighth, per-stage-pair contrasts used limma-trend "
  "on the locked log2-CPM matrix rather than voom precision weights (stated deviation); "
  "the F3 vs F4 pair is low-powered (n = 73/16), as is Fujiwara's F0 vs F1 (n = 12/58). "
  "Ninth, atlas cell types were assigned by canonical markers. Tenth, the locked "
  "Fujiwara matrix is post-normalisation, so genes-detected is degenerate and true "
  "library-size adjustment there rests on an n = 40 raw-count subset. Eleventh, "
  "pathway-specificity verdicts rest on empirical percentiles of 1,000-null "
  "distributions and are accordingly coarse near the threshold (the UCAM suppressor "
  "result sits at p = 0.04995). Twelfth, GSE136103 derives from cirrhotic and healthy "
  "livers, not the MASLD spectrum.")

p("Conflict of Interest", bold=True)
p("The authors declare no conflicts of interest.")
p("Funding", bold=True)
p("This work was supported by [Funding sources to be added].")
p("Data Availability", bold=True)
p("All discovery and validation datasets are publicly available from GEO (GSE135251, "
  "GSE130970, GSE185051; Fujiwara SuperSeries GSE193084 with sub-series GSE192959, "
  "GSE193066, GSE193080) and the Kamzolas et al. (2026) data release; the single-cell "
  "atlas is GSE136103. Analysis code and processed data are available at [repository URL].")

doc.add_heading("References", level=1)
refs_start = find("Angulo, P., Kleiner")
for r in paras[refs_start:]:
    if r.startswith("Supplementary"):
        break
    par = p(r, size=10)
    par.paragraph_format.space_after = Pt(3)
p("Ramachandran, P., Dobie, R., Wilson-Kanamori, J.R., et al., 2020. Resolving the "
  "fibrotic niche of human liver cirrhosis using single-cell transcriptomics. Nature "
  "582, 128–133.", size=10)

doc.add_heading("Supplementary Materials", level=1)
p("Supplementary tables: S1. Full differential-expression table. S2. Complete enrichment "
  "ranking. S3. Benchmark: all panels × cohorts × three splits. S4. Parsimony curve. "
  "S5. Longitudinal per-patient scores. S6. Clinical-subset composition. S7. Network "
  "communities and hubs. S8. NMF cluster-by-stage distributions. S9. Prioritisation. "
  "S10. Fixed operating-point performance. S11. Panel gene lists. S12. Deconvolution: "
  "full correlation table. S13. SPIB post-hoc pairwise comparisons. S14. SPIB "
  "cohort-stratified analysis. S15. Top 20 prioritised transcription factors. S16. "
  "Per-stage-pair DGE tables and combined GSEA results (both cohorts). S17. "
  "Paired-biopsy ferroptosis scores and trajectory tests. S18. Co-expression network "
  "edges, null distribution and modules. S19. Effector-gene donor-pseudobulk values. "
  "S20. Random-set null distributions for every mean-z pathway score (both cohorts; "
  "includes the null-typical pyroptosis and FerrDb-driver rows demoted from Table 7). "
  "S21. Size-matched subsampling distributions. S22. Full 8 × 10 panel-by-NABA-category "
  "matrix (v2023.2.Hs). S23. Basement-membrane gene table. S24. Load manifests and "
  "provenance records.", size=10)
p("Supplementary figures: Fig. S1. Cohort flow diagram. Fig. S2. Deconvolution results. "
  "Fig. S3. Sensitivity analysis. Fig. S4. Per-gene AUROCs. Fig. S5. SPIB "
  "cohort-stratified analysis. Fig. S6. Extra-hepatic pathway enrichment. Fig. S7. "
  "Multi-omic fusion and SHAP interpretability. Fig. S8. Continuous pseudotime "
  "trajectory and longitudinal regression.", size=10)

doc.save(OUT)
print("saved", OUT)
